import streamlit as st
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
import io
import re
import requests
from bs4 import BeautifulSoup

try:
    import pypdf
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False


# ── Page config ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Compliance Reviewer",
    page_icon="📋",
    layout="wide",
)

st.title("📋 Course Compliance Reviewer")
st.caption("Upload a course and a regulation/guidelines doc — get back a Word doc with every conflict and gap highlighted.")

# ── API key setup ─────────────────────────────────────────────────────────────
api_key = None
for k in ["GEMINI_API_KEY", "GROQ_API_KEY"]:
    try:
        val = st.secrets.get(k)
        if val:
            api_key = (k, val)
            break
    except Exception:
        pass

if not api_key:
    st.error("No API key found. Add GEMINI_API_KEY or GROQ_API_KEY to Streamlit secrets.")
    st.stop()


# ── Text extraction helpers ───────────────────────────────────────────────────
def extract_text_from_file(f) -> str:
    name = f.name.lower()
    if name.endswith(".txt"):
        return f.read().decode("utf-8", errors="ignore")
    if name.endswith(".docx"):
        doc = Document(io.BytesIO(f.read()))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if name.endswith(".pdf") and HAS_PYPDF:
        reader = pypdf.PdfReader(io.BytesIO(f.read()))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    return f.read().decode("utf-8", errors="ignore")


def extract_text_from_url(url: str) -> str:
    try:
        resp = requests.get(url, timeout=15, headers={"User-Agent": "Mozilla/5.0"})
        soup = BeautifulSoup(resp.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)[:40_000]
    except Exception as e:
        return f"[Could not fetch {url}: {e}]"


# ── AI review via Hugging Face (no API key needed) ────────────────────────────
SYSTEM_PROMPT = """You are a compliance analyst. You will be given a course script and regulation/guideline documents.

Find every place where the course CONFLICTS with or is MISSING required content from the regulations.

Return a JSON object with this exact structure:
{
  "course_title": "inferred title or 'Untitled Course'",
  "summary": "2-3 sentence executive summary of overall compliance status",
  "findings": [
    {
      "id": 1,
      "type": "ISSUE",
      "section": "section or topic name from the course",
      "quote": "exact verbatim text from the course that is the problem (null for GAP type)",
      "explanation": "clear explanation of the conflict or gap and what the regulation requires",
      "regulation_ref": "which regulation document and what it says"
    }
  ]
}

Rules:
- ISSUE = course text directly conflicts with or contradicts the regulation
- GAP = the regulation requires something completely absent from the course
- Be specific, cite exact quotes for ISSUEs
- Order findings by severity (most critical first)
- Return ONLY valid JSON, no markdown fences, no extra text"""


def run_review(course_text: str, regulation_text: str) -> dict:
    prompt = f"""{SYSTEM_PROMPT}

--- COURSE SCRIPT ---
{course_text[:20000]}

--- REGULATIONS / GUIDELINES ---
{regulation_text[:20000]}

Return ONLY valid JSON:"""

    key_name, key_val = api_key

    if key_name == "GEMINI_API_KEY":
        # Direct REST call to Gemini v1 (not v1beta)
        url = f"https://generativelanguage.googleapis.com/v1/models/gemini-2.5-flash:generateContent?key={key_val}"
        body = {
            "contents": [{"parts": [{"text": prompt}]}],
            "generationConfig": {"temperature": 0.1, "maxOutputTokens": 4000}
        }
        resp = requests.post(url, json=body, timeout=60)
        resp.raise_for_status()
        raw = resp.json()["candidates"][0]["content"]["parts"][0]["text"].strip()

    elif key_name == "GROQ_API_KEY":
        resp = requests.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={"Authorization": f"Bearer {key_val}", "Content-Type": "application/json"},
            json={
                "model": "llama-3.3-70b-versatile",
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.1,
                "max_tokens": 4000,
            },
            timeout=60,
        )
        resp.raise_for_status()
        raw = resp.json()["choices"][0]["message"]["content"].strip()

    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    json_match = re.search(r'\{.*\}', raw, re.DOTALL)
    if json_match:
        raw = json_match.group(0)
    return json.loads(raw)


# ── Word doc builder ──────────────────────────────────────────────────────────
def set_highlight(run, color: str):
    rpr = run._r.get_or_add_rPr()
    hl = OxmlElement("w:highlight")
    hl.set(qn("w:val"), color)
    rpr.append(hl)


def add_heading(doc, text, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    return p


def build_docx(course_text: str, result: dict) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    # ── Header ──
    p = doc.add_paragraph()
    r = p.add_run(result.get("course_title", "Course Review"))
    r.bold = True
    r.font.size = Pt(16)

    p = doc.add_paragraph()
    r = p.add_run("Compliance Review")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0, 70, 127)

    findings = result.get("findings", [])
    issues = [f for f in findings if f.get("type") == "ISSUE"]
    gaps = [f for f in findings if f.get("type") == "GAP"]

    doc.add_paragraph(f"{len(issues)} issue(s) found  •  {len(gaps)} gap(s) found  •  {len(findings)} total findings")
    doc.add_paragraph(result.get("summary", ""))

    # ── Legend ──
    legend = doc.add_paragraph()
    legend.paragraph_format.space_before = Pt(8)
    legend.add_run("Legend:   ")
    yr = legend.add_run("  ISSUE  ")
    set_highlight(yr, "yellow")
    legend.add_run("  = conflicts with regulation        ")
    cr = legend.add_run("  GAP  ")
    set_highlight(cr, "cyan")
    legend.add_run("  = content missing from course")

    # ── Findings summary table ──
    add_heading(doc, "FINDINGS SUMMARY", 12)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, txt in zip(hdr, ["#", "Type", "Summary"]):
        cell.text = txt
        cell.paragraphs[0].runs[0].bold = True

    for f in findings:
        row = table.add_row().cells
        row[0].text = str(f.get("id", ""))
        row[1].text = f.get("type", "")
        row[2].text = f"{f.get('section','')}: {f.get('explanation','')[:120]}..."

    # ── Full script with annotations ──
    doc.add_page_break()
    add_heading(doc, "FULL SCRIPT WITH ANNOTATIONS", 13)

    issue_map = {
        f["quote"]: f
        for f in findings
        if f.get("type") == "ISSUE" and f.get("quote")
    }
    gap_sections = {f.get("section", ""): f for f in findings if f.get("type") == "GAP"}

    for line in course_text.split("\n"):
        if not line.strip():
            continue  # skip blank lines

        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        remaining = line
        matched = False

        for quote, finding in issue_map.items():
            if quote in remaining:
                before, _, after = remaining.partition(quote)
                if before:
                    para.add_run(before)
                yr = para.add_run(quote)
                set_highlight(yr, "yellow")
                marker = para.add_run(f" [{finding['id']}]")
                marker.font.color.rgb = RGBColor(192, 0, 0)
                marker.bold = True
                if after:
                    para.add_run(after)
                matched = True
                break

        if not matched:
            para.add_run(remaining)

        for section, finding in list(gap_sections.items()):
            if section and section.lower() in line.lower():
                gap_para = doc.add_paragraph()
                gap_para.paragraph_format.space_after = Pt(4)
                gr = gap_para.add_run(f"[{finding['id']}] GAP: {finding['explanation']}")
                set_highlight(gr, "cyan")
                del gap_sections[section]

    for finding in gap_sections.values():
        gap_para = doc.add_paragraph()
        gr = gap_para.add_run(f"[{finding['id']}] GAP: {finding['explanation']}")
        set_highlight(gr, "cyan")

    # ── Comment legend ──
    doc.add_page_break()
    add_heading(doc, "Comment Legend", 14)

    for f in findings:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        color = RGBColor(192, 0, 0) if f["type"] == "ISSUE" else RGBColor(0, 70, 127)
        r = p.add_run(f"[{f['id']}] {f['type']}: ")
        r.font.color.rgb = color
        r.bold = True
        r2 = p.add_run(f"{f.get('section', '')} — {f.get('explanation', '')}")
        r2.font.color.rgb = color
        if f.get("regulation_ref"):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Pt(20)
            p2.paragraph_format.space_after = Pt(2)
            r3 = p2.add_run(f"Regulation: {f['regulation_ref']}")
            r3.font.color.rgb = RGBColor(80, 80, 80)
            r3.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── UI ────────────────────────────────────────────────────────────────────────
col1, col2 = st.columns(2)

with col1:
    st.subheader("Course Content")
    course_sku = st.text_input("Course SKU / Title (optional)", placeholder="e.g. RVCT-882")
    course_file = st.file_uploader(
        "Upload course document", type=["docx", "pdf", "txt"], key="course"
    )
    course_text_input = st.text_area(
        "Or paste course text here", height=200, placeholder="Paste transcript or script..."
    )

with col2:
    st.subheader("Regulations / Guidelines")
    reg_files = st.file_uploader(
        "Upload regulation documents", type=["docx", "pdf", "txt"],
        accept_multiple_files=True, key="regs"
    )
    reg_urls = st.text_area(
        "And/or paste URLs to scan (one per line)",
        height=100,
        placeholder="https://cpr.heart.org/...\nhttps://..."
    )

st.divider()
run_btn = st.button("🔍 Run Compliance Review", type="primary", use_container_width=True)

if run_btn:
    course_text = ""
    if course_file:
        course_text = extract_text_from_file(course_file)
    elif course_text_input.strip():
        course_text = course_text_input.strip()

    if not course_text:
        st.error("Please upload a course document or paste the course text.")
        st.stop()

    if course_sku:
        course_text = f"Course SKU: {course_sku}\n\n{course_text}"

    reg_parts = []
    for rf in (reg_files or []):
        reg_parts.append(f"[Document: {rf.name}]\n{extract_text_from_file(rf)}")

    for url in (reg_urls or "").splitlines():
        url = url.strip()
        if url:
            with st.spinner(f"Fetching {url}..."):
                reg_parts.append(f"[URL: {url}]\n{extract_text_from_url(url)}")

    if not reg_parts:
        st.error("Please upload at least one regulation document or enter a URL.")
        st.stop()

    regulation_text = "\n\n".join(reg_parts)

    with st.spinner("Reviewing — this takes 30–60 seconds..."):
        try:
            result = run_review(course_text, regulation_text)
        except Exception as e:
            st.error(f"Review failed: {e}")
            st.stop()

    findings = result.get("findings", [])
    issues = [f for f in findings if f["type"] == "ISSUE"]
    gaps = [f for f in findings if f["type"] == "GAP"]

    st.success(f"Found **{len(issues)} issues** and **{len(gaps)} gaps**")
    st.write(result.get("summary", ""))

    with st.expander("View all findings", expanded=True):
        for f in findings:
            icon = "🟡" if f["type"] == "ISSUE" else "🔵"
            st.markdown(f"**{icon} [{f['id']}] {f['type']} — {f.get('section', '')}**")
            if f.get("quote"):
                st.markdown(f"> {f['quote']}")
            st.markdown(f.get("explanation", ""))
            if f.get("regulation_ref"):
                st.caption(f"Regulation: {f['regulation_ref']}")
            st.divider()

    with st.spinner("Building Word document..."):
        docx_bytes = build_docx(course_text, result)

    title = result.get("course_title", course_sku or "Course").replace(" ", "_")
    st.download_button(
        label="⬇️ Download Reviewed Word Doc",
        data=docx_bytes,
        file_name=f"{title}_Compliance_Review.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary",
        use_container_width=True,
    )
